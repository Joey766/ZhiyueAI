import io
import json
import re

from docx import Document
from pypdf import PdfReader

from ai.ollama_client import OllamaClient, _json_objects
from prompts.resume_parse_prompt import RESUME_PARSE_SCHEMA, build_resume_parse_messages


def extract_resume_text(file_bytes, filename):
    suffix = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    try:
        if suffix == "pdf":
            text = "\n".join(page.extract_text() or "" for page in PdfReader(io.BytesIO(file_bytes)).pages)
        elif suffix == "docx":
            document = Document(io.BytesIO(file_bytes))
            paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            tables = [" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()) for table in document.tables for row in table.rows]
            text = "\n".join(paragraphs + [row for row in tables if row])
        else:
            return {"ok": False, "error": "unsupported"}
    except Exception:
        return {"ok": False, "error": "read_failed"}
    if len(re.sub(r"\s+", "", text)) < 50:
        return {"ok": False, "error": "insufficient_text"}
    return {"ok": True, "text": text}


def _strings(item, keys):
    return {key: str(item.get(key, "") or "").strip() for key in keys}


def validate_resume_result(content):
    try:
        result = json.loads(content)
    except (TypeError, json.JSONDecodeError):
        result = next(_json_objects(content), None) if isinstance(content, str) else None
    if not isinstance(result, dict):
        return None
    basic = result.get("basic")
    if not isinstance(basic, dict):
        return None
    normalized = {
        "basic": _strings(basic, ["name", "english_name", "email", "phone", "location", "linkedin", "portfolio"]),
        "education": [], "experience": [], "projects": [],
        "skills": [str(skill).strip() for skill in result.get("skills", []) if str(skill).strip()],
    }
    groups = {
        "education": ["school", "major", "degree", "start_date", "graduation_date", "gpa"],
        "experience": ["company", "role", "start_date", "end_date", "content"],
        "projects": ["name", "role", "time", "description", "skills"],
    }
    for group, keys in groups.items():
        values = result.get(group, [])
        if not isinstance(values, list): return None
        normalized[group] = [_strings(item, keys) for item in values if isinstance(item, dict)]
    return normalized


def parse_resume(file_bytes, filename):
    extracted = extract_resume_text(file_bytes, filename)
    if not extracted["ok"]:
        return extracted
    response = OllamaClient().chat_json(build_resume_parse_messages(extracted["text"]), schema=RESUME_PARSE_SCHEMA)
    if not response["ok"]:
        return response
    parsed = validate_resume_result(response["content"])
    if parsed is None:
        return {"ok": False, "error": "invalid_json"}
    return {"ok": True, "result": parsed, "text_length": len(extracted["text"])}
